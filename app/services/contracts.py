from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import io
from datetime import datetime, date, timedelta
from typing import Optional, List
from app.logger import get_logger

logger = get_logger(__name__)

# ── Design constants ───────────────────────────────────
TEAL = RGBColor(0x2D, 0x7A, 0x6B)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x5A, 0x55, 0x50)


def _style_heading(para, color=TEAL):
    for run in para.runs:
        run.font.color.rgb = color
        run.font.bold = True


def _add_kv_table(doc, rows: list) -> None:
    """Add a clean 2-column key-value table"""
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(value) if value else "N/A"
        run = row[0].paragraphs[0].runs
        if run:
            run[0].bold = True
            run[0].font.color.rgb = MUTED
    doc.add_paragraph()


def generate_contract(
    client: dict,
    contract_data: dict
) -> bytes:
    """Generate a professional service contract"""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Title
    title = doc.add_heading("SERVICE AGREEMENT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _style_heading(title, TEAL)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"Contract No: "
        f"{contract_data.get('contract_number', 'N/A')}\n"
        f"Date: {datetime.now().strftime('%B %d, %Y')}\n"
        f"Valid Until: "
        f"{contract_data.get('valid_until', 'N/A')}"
    )
    run.italic = True
    run.font.color.rgb = MUTED
    doc.add_paragraph()

    # Parties
    h = doc.add_heading("1. PARTIES", level=1)
    _style_heading(h)
    _add_kv_table(doc, [
        ("Service Provider",
         contract_data.get("provider_name", "Your Company")),
        ("Provider Address",
         contract_data.get("provider_address", "")),
        ("Provider Email",
         contract_data.get("provider_email", "")),
        ("Client Name", client.get("name", "")),
        ("Client Company", client.get("company", "")),
        ("Client Email", client.get("email", "")),
        ("Client Phone", client.get("phone", ""))
    ])

    # Scope
    h = doc.add_heading("2. SCOPE OF WORK", level=1)
    _style_heading(h)
    doc.add_paragraph(
        contract_data.get(
            "scope_of_work",
            f"Provider will deliver: "
            f"{client.get('service', 'services as discussed.')}"
        )
    )
    doc.add_paragraph()

    # Deliverables
    if contract_data.get("deliverables"):
        h = doc.add_heading("3. DELIVERABLES", level=1)
        _style_heading(h)
        for i, d in enumerate(
            contract_data["deliverables"], 1
        ):
            doc.add_paragraph(
                f"{i}. {d}",
                style="List Number"
            )
        doc.add_paragraph()

    # Timeline
    if contract_data.get("timeline"):
        h = doc.add_heading("4. TIMELINE", level=1)
        _style_heading(h)
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        hdr = table.add_row().cells
        hdr[0].text = "Milestone"
        hdr[1].text = "Target Date"
        for run in hdr[0].paragraphs[0].runs:
            run.bold = True
        for run in hdr[1].paragraphs[0].runs:
            run.bold = True
        for m in contract_data["timeline"]:
            row = table.add_row().cells
            row[0].text = m.get("milestone", "")
            row[1].text = m.get("date", "")
        doc.add_paragraph()

    # Payment
    h = doc.add_heading("5. PAYMENT TERMS", level=1)
    _style_heading(h)
    _add_kv_table(doc, [
        ("Total Amount",
         f"${contract_data.get('total_amount', '0.00')}"),
        ("Payment Schedule",
         contract_data.get("payment_schedule", "Net 30")),
        ("Payment Method",
         contract_data.get("payment_method", "Bank Transfer")),
        ("Late Fee",
         contract_data.get("late_fee", "1.5% per month"))
    ])

    # Terms
    h = doc.add_heading("6. TERMS AND CONDITIONS", level=1)
    _style_heading(h)
    terms = contract_data.get("terms", [
        "Provider will maintain confidentiality of all "
        "client information.",
        "Client will provide timely feedback and approvals.",
        "Either party may terminate with 30 days written "
        "notice.",
        "Provider retains ownership of work until full "
        "payment is received.",
        "This agreement is governed by the laws of Utah, USA.",
        "Disputes will be resolved through mediation.",
        "Force majeure clause applies.",
        "Amendments must be in writing and signed by both "
        "parties."
    ])
    for i, term in enumerate(terms, 1):
        doc.add_paragraph(f"{i}. {term}")
    doc.add_paragraph()

    # Confidentiality
    h = doc.add_heading("7. CONFIDENTIALITY", level=1)
    _style_heading(h)
    doc.add_paragraph(
        contract_data.get(
            "confidentiality",
            "Both parties agree to keep all shared information "
            "strictly confidential and not disclose it to "
            "third parties without prior written consent."
        )
    )
    doc.add_paragraph()

    # IP
    h = doc.add_heading("8. INTELLECTUAL PROPERTY", level=1)
    _style_heading(h)
    doc.add_paragraph(
        contract_data.get(
            "ip_clause",
            "Upon receipt of full payment, all deliverables "
            "become the exclusive property of the Client. "
            "Provider may use the work in portfolio with "
            "Client approval."
        )
    )
    doc.add_paragraph()

    # Signatures
    h = doc.add_heading("SIGNATURES", level=1)
    _style_heading(h)
    sig_table = doc.add_table(rows=0, cols=2)
    sig_table.style = "Table Grid"
    hdr = sig_table.add_row().cells
    hdr[0].text = "Service Provider"
    hdr[1].text = "Client"
    for c in [hdr[0], hdr[1]]:
        for run in c.paragraphs[0].runs:
            run.bold = True

    row = sig_table.add_row().cells
    row[0].text = (
        f"\n\nSignature: _______________________\n"
        f"Name: {contract_data.get('provider_name', '')}\n"
        f"Date: _______________________"
    )
    row[1].text = (
        f"\n\nSignature: _______________________\n"
        f"Name: {client.get('name', '')}\n"
        f"Date: _______________________"
    )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    logger.info(
        f"Contract generated: {client.get('client_id')}"
    )
    return buffer.getvalue()


def generate_invoice(
    client: dict,
    invoice_data: dict
) -> bytes:
    """Generate a professional invoice"""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    title = doc.add_heading("INVOICE", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _style_heading(title, TEAL)
    doc.add_paragraph()

    invoice_no = invoice_data.get(
        "invoice_number",
        f"INV-{datetime.now().strftime('%Y%m%d%H%M')}"
    )
    due = invoice_data.get(
        "due_date",
        (date.today() + timedelta(days=30))
        .strftime("%B %d, %Y")
    )

    _add_kv_table(doc, [
        ("Invoice No", invoice_no),
        ("Date", datetime.now().strftime("%B %d, %Y")),
        ("Due Date", due),
        ("Status", invoice_data.get("status", "Unpaid"))
    ])

    h = doc.add_heading("BILLING DETAILS", level=1)
    _style_heading(h)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    hdr = table.add_row().cells
    hdr[0].text = "From"
    hdr[1].text = "Bill To"
    for c in [hdr[0], hdr[1]]:
        for run in c.paragraphs[0].runs:
            run.bold = True

    row = table.add_row().cells
    row[0].text = (
        f"{invoice_data.get('provider_name', 'Your Company')}\n"
        f"{invoice_data.get('provider_email', '')}\n"
        f"{invoice_data.get('provider_address', '')}"
    )
    row[1].text = (
        f"{client.get('name', '')}\n"
        f"{client.get('company', '')}\n"
        f"{client.get('email', '')}\n"
        f"{client.get('phone', '')}"
    )
    doc.add_paragraph()

    h = doc.add_heading("SERVICES", level=1)
    _style_heading(h)
    items_table = doc.add_table(rows=0, cols=4)
    items_table.style = "Table Grid"
    hdr = items_table.add_row().cells
    for i, h_text in enumerate(
        ["Description", "Qty", "Unit Price", "Total"]
    ):
        hdr[i].text = h_text
        if hdr[i].paragraphs[0].runs:
            hdr[i].paragraphs[0].runs[0].bold = True

    subtotal = 0.0
    for item in invoice_data.get("line_items", []):
        qty = float(item.get("quantity", 1))
        price = float(item.get("unit_price", 0))
        total = qty * price
        subtotal += total
        row = items_table.add_row().cells
        row[0].text = item.get("description", "")
        row[1].text = str(int(qty) if qty == int(qty) else qty)
        row[2].text = f"${price:,.2f}"
        row[3].text = f"${total:,.2f}"

    doc.add_paragraph()

    tax_rate = float(invoice_data.get("tax_rate", 0))
    tax_amt = subtotal * (tax_rate / 100)
    discount = float(invoice_data.get("discount", 0))
    total_due = subtotal + tax_amt - discount

    totals_table = doc.add_table(rows=0, cols=2)
    totals_table.style = "Table Grid"
    for label, value in [
        ("Subtotal", f"${subtotal:,.2f}"),
        (f"Tax ({tax_rate}%)", f"${tax_amt:,.2f}"),
        ("Discount", f"-${discount:,.2f}"),
        ("TOTAL DUE", f"${total_due:,.2f}")
    ]:
        row = totals_table.add_row().cells
        row[0].text = label
        row[1].text = value
        if label == "TOTAL DUE":
            for cell in [row[0], row[1]]:
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].bold = True
                    cell.paragraphs[0].runs[0].font.color.rgb = TEAL

    doc.add_paragraph()
    h = doc.add_heading("PAYMENT INSTRUCTIONS", level=1)
    _style_heading(h)
    doc.add_paragraph(
        invoice_data.get(
            "payment_instructions",
            "Please make payment via bank transfer. "
            "Reference the invoice number in your payment."
        )
    )

    if invoice_data.get("notes"):
        h = doc.add_heading("NOTES", level=1)
        _style_heading(h)
        doc.add_paragraph(invoice_data["notes"])

    doc.add_paragraph()
    footer = doc.add_paragraph("Thank you for your business.")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if footer.runs:
        footer.runs[0].italic = True
        footer.runs[0].font.color.rgb = MUTED

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    logger.info(
        f"Invoice generated: {client.get('client_id')}"
    )
    return buffer.getvalue()


def generate_proposal(
    client: dict,
    proposal_data: dict
) -> bytes:
    """Generate a professional business proposal"""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    title = doc.add_heading("BUSINESS PROPOSAL", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _style_heading(title, TEAL)

    sub = doc.add_paragraph(
        f"Prepared for: "
        f"{client.get('company') or client.get('name', '')}\n"
        f"Prepared by: "
        f"{proposal_data.get('provider_name', 'Your Company')}\n"
        f"Date: {datetime.now().strftime('%B %d, %Y')}\n"
        f"Valid Until: "
        f"{proposal_data.get('valid_until', '30 days')}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if sub.runs:
        sub.runs[0].italic = True
        sub.runs[0].font.color.rgb = MUTED
    doc.add_paragraph()

    sections_content = [
        ("1. EXECUTIVE SUMMARY",
         proposal_data.get(
             "executive_summary",
             f"We are pleased to present this proposal for "
             f"{client.get('service', 'our services')} to "
             f"{client.get('company') or client.get('name')}."
         )),
        ("2. PROBLEM STATEMENT",
         proposal_data.get(
             "problem_statement",
             client.get("notes", "As discussed.")
         )),
        ("3. PROPOSED SOLUTION",
         proposal_data.get(
             "proposed_solution",
             f"We propose to deliver "
             f"{client.get('service', 'our services')} "
             f"tailored to your specific needs."
         ))
    ]

    for heading_text, content in sections_content:
        h = doc.add_heading(heading_text, level=1)
        _style_heading(h)
        doc.add_paragraph(content)
        doc.add_paragraph()

    if proposal_data.get("scope_items"):
        h = doc.add_heading("4. SCOPE OF WORK", level=1)
        _style_heading(h)
        for item in proposal_data["scope_items"]:
            doc.add_paragraph(f"- {item}")
        doc.add_paragraph()

    if proposal_data.get("timeline"):
        h = doc.add_heading("5. PROJECT TIMELINE", level=1)
        _style_heading(h)
        tl_table = doc.add_table(rows=0, cols=3)
        tl_table.style = "Table Grid"
        hdr = tl_table.add_row().cells
        for i, txt in enumerate(
            ["Phase", "Description", "Duration"]
        ):
            hdr[i].text = txt
            if hdr[i].paragraphs[0].runs:
                hdr[i].paragraphs[0].runs[0].bold = True
        for phase in proposal_data["timeline"]:
            row = tl_table.add_row().cells
            row[0].text = phase.get("phase", "")
            row[1].text = phase.get("description", "")
            row[2].text = phase.get("duration", "")
        doc.add_paragraph()

    h = doc.add_heading("6. INVESTMENT", level=1)
    _style_heading(h)
    if proposal_data.get("pricing_tiers"):
        for tier in proposal_data["pricing_tiers"]:
            doc.add_heading(
                tier.get("name", "Option"), level=2
            )
            doc.add_paragraph(
                f"Price: ${tier.get('price', 'TBD')}\n"
                f"{tier.get('description', '')}"
            )
            if tier.get("includes"):
                for item in tier["includes"]:
                    doc.add_paragraph(f"  + {item}")
    else:
        doc.add_paragraph(
            f"Total Investment: "
            f"${proposal_data.get('total_price', 'TBD')}\n"
            f"{proposal_data.get('pricing_notes', '')}"
        )
    doc.add_paragraph()

    if proposal_data.get("why_us"):
        h = doc.add_heading("7. WHY CHOOSE US", level=1)
        _style_heading(h)
        for point in proposal_data["why_us"]:
            doc.add_paragraph(f"+ {point}")
        doc.add_paragraph()

    h = doc.add_heading("8. NEXT STEPS", level=1)
    _style_heading(h)
    for i, step in enumerate(proposal_data.get(
        "next_steps",
        ["Review proposal", "Schedule call",
         "Sign agreement", "Project kickoff"]
    ), 1):
        doc.add_paragraph(f"{i}. {step}")
    doc.add_paragraph()

    cta = doc.add_paragraph(
        proposal_data.get(
            "call_to_action",
            f"Ready to get started? Contact us at "
            f"{proposal_data.get('provider_email', '')}"
        )
    )
    cta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cta.runs:
        cta.runs[0].bold = True
        cta.runs[0].font.color.rgb = TEAL

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    logger.info(
        f"Proposal generated: {client.get('client_id')}"
    )
    return buffer.getvalue()