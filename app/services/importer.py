import io
import csv
import uuid
import asyncio
from datetime import datetime
from typing import List, Dict, Tuple, AsyncGenerator
from app.logger import get_logger

logger = get_logger(__name__)

VALID_PRIORITIES = {"High", "Medium", "Low"}
VALID_STAGES = {
    "New", "Contacted",
    "Consultation Scheduled",
    "Proposal Sent", "Won", "Lost"
}

COL_MAP = {
    "name": [
        "name", "full name", "client name",
        "contact name", "contact"
    ],
    "email": [
        "email", "email address", "e-mail", "mail"
    ],
    "company": [
        "company", "organization", "org",
        "business", "company name"
    ],
    "phone": [
        "phone", "phone number", "mobile",
        "cell", "telephone", "contact number"
    ],
    "service": [
        "service", "service type", "product",
        "interest", "needs"
    ],
    "priority": [
        "priority", "priority level", "importance"
    ],
    "stage": [
        "stage", "pipeline stage", "status"
    ],
    "notes": [
        "notes", "note", "comments", "comment",
        "description", "details"
    ]
}


def _map_columns(headers: List[str]) -> Dict[str, int]:
    """Map file headers to our field names"""
    mapping = {}
    headers_lower = [h.strip().lower() for h in headers]
    for field, variants in COL_MAP.items():
        for i, h in enumerate(headers_lower):
            if h in variants:
                mapping[field] = i
                break
    return mapping


def _clean_row(
    row: List[str],
    col_map: Dict[str, int]
) -> Dict:
    """Extract and normalise a single row"""

    def get(field) -> str | None:
        idx = col_map.get(field)
        if idx is None or idx >= len(row):
            return None
        val = str(row[idx]).strip()
        # Treat nan, None, empty as missing
        if not val or val.lower() in ("nan", "none", "null"):
            return None
        return val

    priority = get("priority")
    if priority and priority.capitalize() in VALID_PRIORITIES:
        priority = priority.capitalize()
    else:
        priority = "Medium"

    stage = get("stage")
    if stage:
        matched = next(
            (vs for vs in VALID_STAGES
             if stage.lower() == vs.lower()),
            None
        )
        stage = matched if matched else "New"
    else:
        stage = "New"

    return {
        "name": get("name"),
        "email": get("email"),
        "company": get("company"),
        "phone": get("phone"),
        "service": get("service"),
        "priority": priority,
        "stage": stage,
        "notes": get("notes")
    }


def _validate_row(
    row_data: Dict,
    row_num: int
) -> Tuple[bool, str]:
    """Validate a single row. Returns (valid, error_msg)"""
    if not row_data.get("name"):
        return False, f"Row {row_num}: Missing required field 'Name'"

    email = row_data.get("email")
    if email and "@" not in email:
        return False, (
            f"Row {row_num}: Invalid email '{email}'"
        )
    if email and len(email) > 254:
        return False, (
            f"Row {row_num}: Email too long"
        )

    return True, ""


def _find_header_row(
    raw_rows: List[List[str]]
) -> Tuple[List[str] | None, int]:
    """Find the first non-empty row to use as header"""
    for i, row in enumerate(raw_rows):
        if any(str(cell).strip() for cell in row):
            return row, i + 1
    return None, 0


def _process_rows(
    raw_rows: List[List[str]]
) -> Tuple[List[Dict], List[str], List[str]]:
    """
    Process raw rows into validated client dicts.
    Returns (valid_rows, detected_columns, errors)
    """
    header_row, data_start = _find_header_row(raw_rows)

    if header_row is None:
        raise ValueError("File appears to be empty")

    col_map = _map_columns(header_row)

    if "name" not in col_map:
        # Try to be helpful about what was found
        found = [str(h).strip() for h in header_row if str(h).strip()]
        raise ValueError(
            f"Could not find a 'Name' column. "
            f"Found columns: {', '.join(found[:8])}. "
            f"Please add a 'Name' header to your file."
        )

    detected_cols = list(col_map.keys())
    valid_rows = []
    errors = []
    seen_emails: set = set()

    for i, row in enumerate(
        raw_rows[data_start:],
        start=data_start + 1
    ):
        # Skip fully empty rows
        if not any(str(cell).strip() for cell in row):
            continue

        row_data = _clean_row(row, col_map)
        is_valid, err = _validate_row(row_data, i)

        if not is_valid:
            errors.append(err)
            continue

        email = row_data.get("email")
        if email:
            if email.lower() in seen_emails:
                errors.append(
                    f"Row {i}: Duplicate email "
                    f"'{email}' within file — skipped"
                )
                continue
            seen_emails.add(email.lower())

        valid_rows.append(row_data)

    return valid_rows, detected_cols, errors


def parse_excel(
    file_bytes: bytes
) -> Tuple[List[Dict], List[str], List[str]]:
    """Parse an Excel (.xlsx/.xls) file"""
    try:
        import openpyxl
    except ImportError:
        raise ValueError(
            "openpyxl is not installed. "
            "Add it to requirements.txt."
        )

    try:
        wb = openpyxl.load_workbook(
            io.BytesIO(file_bytes),
            read_only=True,
            data_only=True
        )
        ws = wb.active
        rows = []
        for row in ws.iter_rows(values_only=True):
            rows.append([
                str(cell) if cell is not None else ""
                for cell in row
            ])
        wb.close()
    except Exception as e:
        raise ValueError(f"Could not read Excel file: {e}")

    if not rows:
        raise ValueError("Excel file is empty")

    return _process_rows(rows)


def parse_csv(
    file_bytes: bytes
) -> Tuple[List[Dict], List[str], List[str]]:
    """Parse a CSV file with encoding detection"""
    for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            text = file_bytes.decode(encoding)
            break
        except (UnicodeDecodeError, Exception):
            continue
    else:
        raise ValueError(
            "Could not decode CSV file. "
            "Please save it as UTF-8."
        )

    try:
        # Detect delimiter
        sample = text[:2048]
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
        reader = csv.reader(io.StringIO(text), dialect)
    except csv.Error:
        reader = csv.reader(io.StringIO(text))

    rows = list(reader)
    if not rows:
        raise ValueError("CSV file is empty")

    return _process_rows(rows)


# ── Sync batch insert ──────────────────────────────────

def _batch_insert_sync(
    rows_to_insert: List[Dict],
    spreadsheet
) -> Tuple[List[Dict], List[Dict]]:
    """
    Insert all rows in a SINGLE Google Sheets API call.
    Must be called from a thread pool — this is blocking.
    """
    ws = spreadsheet.worksheet("Clients")
    headers = ws.row_values(1)
    col = {h: i for i, h in enumerate(headers)}
    now = datetime.utcnow().isoformat()

    created = []
    failed = []
    batch_rows = []

    for row in rows_to_insert:
        try:
            client_id = "CL-" + str(uuid.uuid4())[:8].upper()
            sheet_row = [""] * len(headers)

            def set_col(
                field: str,
                value,
                _row=sheet_row,
                _col=col
            ):
                if field in _col and value is not None:
                    _row[_col[field]] = str(value)

            set_col("Client ID", client_id)
            set_col("Created At", now)
            set_col("Name", row.get("name", ""))
            set_col("Company", row.get("company", ""))
            set_col("Email", row.get("email", ""))
            set_col("Phone", row.get("phone", ""))
            set_col("Service", row.get("service", ""))
            set_col("Priority",
                    row.get("priority", "Medium"))
            set_col("Stage", row.get("stage", "New"))
            set_col("Notes", row.get("notes", ""))
            set_col("Archived", "No")
            set_col("Folder URL", "")
            set_col("Report URL", "")
            set_col("Next Follow-up", "")

            batch_rows.append(sheet_row)
            created.append({
                "name": row.get("name"),
                "client_id": client_id,
                "email": row.get("email")
            })

        except Exception as e:
            failed.append({
                "name": row.get("name", "Unknown"),
                "error": str(e)
            })

    if batch_rows:
        ws.append_rows(
            batch_rows,
            value_input_option="RAW",
            insert_data_option="INSERT_ROWS",
            table_range="A1"
        )

    return created, failed


# ── Async streaming import ─────────────────────────────

async def bulk_import_clients_stream(
    rows: List[Dict],
    sheets_module,
    check_duplicates: bool = True
) -> AsyncGenerator[Dict, None]:
    """
    Import clients with live Server-Sent Events progress.
    Yields status dicts compatible with SSE format.
    """
    total = len(rows)

    yield {
        "type": "start",
        "total": total,
        "message": f"Starting import of {total} clients"
    }
    await asyncio.sleep(0)

    # Step 1: Fetch existing emails for dedup
    existing_emails: set = set()
    if check_duplicates:
        yield {
            "type": "progress",
            "step": "checking",
            "message": "Checking for duplicates in CRM...",
            "percent": 5
        }
        await asyncio.sleep(0)

        try:
            existing = await asyncio.to_thread(
                sheets_module.get_all_clients_sync
            )
            existing_emails = {
                c.get("email", "").lower()
                for c in existing
                if c.get("email")
            }
        except Exception as e:
            logger.warning(
                f"Duplicate check failed: {e}"
            )

    # Step 2: Filter
    rows_to_insert: List[Dict] = []
    skipped: List[Dict] = []
    seen_in_batch: set = set()

    for row in rows:
        email = (row.get("email") or "").strip().lower()
        if check_duplicates and email:
            if email in existing_emails:
                skipped.append({
                    "name": row.get("name"),
                    "email": row.get("email"),
                    "reason": "Email already exists in CRM"
                })
                continue
            if email in seen_in_batch:
                skipped.append({
                    "name": row.get("name"),
                    "email": row.get("email"),
                    "reason": "Duplicate email within file"
                })
                continue
            seen_in_batch.add(email)
        rows_to_insert.append(row)

    yield {
        "type": "progress",
        "step": "filtered",
        "message": (
            f"{len(rows_to_insert)} to import, "
            f"{len(skipped)} duplicates skipped"
        ),
        "percent": 15,
        "skipped": len(skipped)
    }
    await asyncio.sleep(0)

    if not rows_to_insert:
        yield {
            "type": "complete",
            "imported": 0,
            "skipped": len(skipped),
            "failed": 0,
            "skipped_clients": skipped,
            "imported_clients": [],
            "failed_clients": [],
            "message": "No new clients to import"
        }
        return

    # Step 3: Batch insert in chunks
    CHUNK_SIZE = 25
    chunks = [
        rows_to_insert[i: i + CHUNK_SIZE]
        for i in range(0, len(rows_to_insert), CHUNK_SIZE)
    ]

    all_created: List[Dict] = []
    all_failed: List[Dict] = []

    try:
        spreadsheet = await asyncio.to_thread(
            sheets_module.get_spreadsheet
        )
    except Exception as e:
        yield {
            "type": "error",
            "message": f"Could not connect to Google Sheets: {e}"
        }
        return

    for i, chunk in enumerate(chunks):
        pct = 15 + int((i / len(chunks)) * 80)
        yield {
            "type": "progress",
            "step": "inserting",
            "message": (
                f"Inserting batch {i+1} of {len(chunks)} "
                f"({len(chunk)} clients)"
            ),
            "percent": pct,
            "inserted_so_far": len(all_created)
        }
        await asyncio.sleep(0)

        try:
            created, failed = await asyncio.to_thread(
                _batch_insert_sync, chunk, spreadsheet
            )
        except Exception as e:
            logger.error(f"Batch {i+1} failed: {e}")
            all_failed.extend([
                {"name": r.get("name"), "error": str(e)}
                for r in chunk
            ])
            continue

        all_created.extend(created)
        all_failed.extend(failed)

        for client in created:
            yield {
                "type": "client_added",
                "client": client
            }
            await asyncio.sleep(0)

    # Clear cache so next read is fresh
    try:
        from app.services.sheets import _cache_clear
        _cache_clear("all_records")
    except Exception:
        pass

    yield {
        "type": "complete",
        "imported": len(all_created),
        "skipped": len(skipped),
        "failed": len(all_failed),
        "imported_clients": all_created,
        "skipped_clients": skipped,
        "failed_clients": all_failed,
        "percent": 100,
        "message": (
            f"Complete: {len(all_created)} imported, "
            f"{len(skipped)} skipped, "
            f"{len(all_failed)} failed"
        )
    }


# ── Non-streaming fallback ─────────────────────────────

async def bulk_import_clients(
    rows: List[Dict],
    sheets_module,
    check_duplicates: bool = True
) -> Dict:
    """Non-streaming bulk import for direct API calls"""
    existing_emails: set = set()

    if check_duplicates:
        try:
            existing = await asyncio.to_thread(
                sheets_module.get_all_clients_sync
            )
            existing_emails = {
                c.get("email", "").lower()
                for c in existing
                if c.get("email")
            }
        except Exception as e:
            logger.warning(
                f"Duplicate check failed: {e}"
            )

    rows_to_insert: List[Dict] = []
    skipped: List[Dict] = []
    seen_in_batch: set = set()

    for row in rows:
        email = (row.get("email") or "").strip().lower()
        if check_duplicates and email:
            if email in existing_emails or \
               email in seen_in_batch:
                skipped.append({
                    "name": row.get("name"),
                    "email": row.get("email"),
                    "reason": "Duplicate"
                })
                continue
            seen_in_batch.add(email)
        rows_to_insert.append(row)

    if not rows_to_insert:
        return {
            "total_rows": len(rows),
            "imported": 0,
            "skipped": len(skipped),
            "failed": 0,
            "imported_clients": [],
            "skipped_clients": skipped,
            "failed_clients": []
        }

    spreadsheet = await asyncio.to_thread(
        sheets_module.get_spreadsheet
    )
    created, failed = await asyncio.to_thread(
        _batch_insert_sync, rows_to_insert, spreadsheet
    )

    # Clear cache
    try:
        from app.services.sheets import _cache_clear
        _cache_clear("all_records")
    except Exception:
        pass

    return {
        "total_rows": len(rows),
        "imported": len(created),
        "skipped": len(skipped),
        "failed": len(failed),
        "imported_clients": created,
        "skipped_clients": skipped,
        "failed_clients": failed
    }
    
def extract_word_text(file_bytes: bytes) -> str:
    """
    Extract plain text from a Word (.docx) file.
    Used when user uploads a Word doc to ARIA chat.
    """
    try:
        from docx import Document
        import io
        doc = Document(io.BytesIO(file_bytes))
        lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                )
                if row_text:
                    lines.append(row_text)
        return "\n".join(lines)
    except Exception as e:
        raise ValueError(
            f"Could not read Word document: {e}. "
            f"Ensure it is a valid .docx file."
        )


def detect_file_type(
    filename: str,
    file_bytes: bytes
) -> str:
    """
    Detect file type from filename and magic bytes.
    Returns: 'pdf', 'excel', 'csv', 'word', 'unknown'
    """
    name = filename.lower().strip()

    if name.endswith(".pdf"):
        return "pdf"
    if name.endswith((".xlsx", ".xls")):
        return "excel"
    if name.endswith(".csv"):
        return "csv"
    if name.endswith(".docx"):
        return "word"

    # Fallback: check magic bytes
    if file_bytes[:4] == b"%PDF":
        return "pdf"
    if file_bytes[:2] in (b"PK", ):
        # Could be xlsx or docx — check filename
        if "doc" in name:
            return "word"
        return "excel"

    return "unknown"


async def process_chat_file(
    filename: str,
    file_bytes: bytes,
    user_message: str = ""
) -> dict:
    """
    Process any file uploaded to ARIA chat.
    Returns structured result with type, text, and action.

    Returns:
        {
            "file_type": "pdf|excel|csv|word|unknown",
            "action": "import|analyze|unknown",
            "text": "extracted text if analyze",
            "rows": [...] if import,
            "detected_columns": [...] if import,
            "parse_errors": [...] if import,
            "summary": "what was found"
        }
    """
    file_type = detect_file_type(filename, file_bytes)

    # Determine action from user message
    import_keywords = [
        "import", "bulk", "upload clients",
        "add clients", "load clients", "insert"
    ]
    analyze_keywords = [
        "analyze", "read", "summarize", "what",
        "tell me", "extract", "review", "check"
    ]

    msg_lower = user_message.lower()
    wants_import = any(
        kw in msg_lower for kw in import_keywords
    )
    wants_analyze = any(
        kw in msg_lower for kw in analyze_keywords
    )

    if file_type == "pdf":
        # PDFs always analyzed — extract text
        from app.services.pdf import extract_pdf_text
        extracted = extract_pdf_text(file_bytes)
        return {
            "file_type": "pdf",
            "action": "analyze",
            "text": extracted["raw_text"],
            "page_count": extracted["page_count"],
            "word_count": extracted["word_count"],
            "has_tables": extracted["has_tables"],
            "summary": (
                f"PDF with {extracted['page_count']} pages, "
                f"{extracted['word_count']} words"
            )
        }

    elif file_type in ("excel", "csv"):
        # Excel/CSV: import by default unless user says analyze
        if wants_analyze and not wants_import:
            # Treat as data to analyze
            if file_type == "excel":
                import openpyxl
                import io as _io
                wb = openpyxl.load_workbook(
                    _io.BytesIO(file_bytes),
                    read_only=True,
                    data_only=True
                )
                ws = wb.active
                lines = []
                for row in ws.iter_rows(values_only=True):
                    row_text = " | ".join(
                        str(c) for c in row
                        if c is not None
                    )
                    if row_text.strip():
                        lines.append(row_text)
                wb.close()
                text = "\n".join(lines[:100])
            else:
                try:
                    text = file_bytes.decode("utf-8")
                except UnicodeDecodeError:
                    text = file_bytes.decode("latin-1")
                # Cap to 100 lines
                text = "\n".join(
                    text.split("\n")[:100]
                )

            return {
                "file_type": file_type,
                "action": "analyze",
                "text": text,
                "summary": (
                    f"{file_type.upper()} file ready for analysis"
                )
            }

        else:
            # Default: parse as client import
            if file_type == "excel":
                rows, cols, errors = parse_excel(file_bytes)
            else:
                rows, cols, errors = parse_csv(file_bytes)

            return {
                "file_type": file_type,
                "action": "import",
                "rows": rows,
                "detected_columns": cols,
                "parse_errors": errors,
                "total_valid_rows": len(rows),
                "summary": (
                    f"{len(rows)} valid clients found, "
                    f"{len(errors)} errors"
                )
            }

    elif file_type == "word":
        text = extract_word_text(file_bytes)
        return {
            "file_type": "word",
            "action": "analyze",
            "text": text,
            "word_count": len(text.split()),
            "summary": (
                f"Word document with "
                f"{len(text.split())} words"
            )
        }

    else:
        return {
            "file_type": "unknown",
            "action": "unknown",
            "text": "",
            "summary": (
                f"Unsupported file type: {filename}. "
                f"Supported: PDF, Excel, CSV, Word"
            )
        }